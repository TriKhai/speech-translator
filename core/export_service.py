"""
core/export_service.py — Xuất transcript ra .txt / .srt / .docx

Usage:
    from core.export_service import ExportService
    segments = db.get_segments(session_id)   # list[dict]
    ExportService.export(segments, "/path/to/output.srt")

Segment dict cần có:
    {
        "text_en":       str,
        "text_vi":       str,
        "speaker_index": int,
        "start_time":    float,   # seconds
        "end_time":      float,
    }
"""

import os
from datetime import timedelta


class ExportService:

    @staticmethod
    def export(segments: list[dict], path: str, lang: str = "vi"):
        """
        Auto-detect format từ extension.
        lang: "vi" | "en" | "both"
        """
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            ExportService.to_txt(segments, path, lang)
        elif ext == ".srt":
            ExportService.to_srt(segments, path, lang)
        elif ext in (".docx", ".doc"):
            ExportService.to_docx(segments, path, lang)
        else:
            raise ValueError(f"Unsupported format: {ext}")

    # ── TXT ─────────────────────────────────────────────────────────────────

    @staticmethod
    def to_txt(segments: list[dict], path: str, lang: str = "vi"):
        lines = []
        prev_speaker = None

        for seg in segments:
            spk = seg.get("speaker_index", 1)
            if spk != prev_speaker:
                if lines:
                    lines.append("")
                lines.append(f"[SPEAKER {spk}]")
                prev_speaker = spk

            text = ExportService._get_text(seg, lang)
            ts   = ExportService._fmt_ts_simple(seg["start_time"])
            lines.append(f"[{ts}] {text}")

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"[EXPORT] TXT → {path} ({len(segments)} segments)")

    # ── SRT ─────────────────────────────────────────────────────────────────

    @staticmethod
    def to_srt(segments: list[dict], path: str, lang: str = "vi"):
        """
        Chuẩn SRT. Nếu lang="both": EN ở dòng 1, VI ở dòng 2.
        """
        lines = []
        for i, seg in enumerate(segments, 1):
            start = ExportService._fmt_srt_time(seg["start_time"])
            end   = ExportService._fmt_srt_time(seg["end_time"])
            spk   = seg.get("speaker_index", 1)

            lines.append(str(i))
            lines.append(f"{start} --> {end}")

            if lang == "both":
                lines.append(f"[S{spk}] {seg.get('text_en', '')}")
                lines.append(seg.get("text_vi", ""))
            else:
                text = ExportService._get_text(seg, lang)
                lines.append(f"[S{spk}] {text}")

            lines.append("")   # blank line giữa các block

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"[EXPORT] SRT → {path} ({len(segments)} segments)")

    # ── DOCX ────────────────────────────────────────────────────────────────

    @staticmethod
    def to_docx(segments: list[dict], path: str, lang: str = "vi"):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "[EXPORT] Cần cài python-docx:\n  pip install python-docx"
            )

        doc = Document()

        # Title
        title = doc.add_heading("Transcript", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle — thời lượng
        if segments:
            duration = segments[-1]["end_time"]
            m, s = divmod(int(duration), 60)
            h, m = divmod(m, 60)
            sub = doc.add_paragraph(f"Thời lượng: {h:02d}:{m:02d}:{s:02d}")
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Speaker color map (RGB)
        COLORS = [
            RGBColor(0x60, 0xA5, 0xFA),   # blue
            RGBColor(0x34, 0xD3, 0x99),   # green
            RGBColor(0xF4, 0x72, 0xB6),   # pink
            RGBColor(0xFB, 0xBF, 0x24),   # amber
            RGBColor(0xA7, 0x8B, 0xFA),   # violet
            RGBColor(0x38, 0xBD, 0xF8),   # sky
        ]

        prev_spk = None
        for seg in segments:
            spk  = seg.get("speaker_index", 1)
            color = COLORS[(spk - 1) % len(COLORS)]

            # Speaker header (chỉ khi đổi speaker)
            if spk != prev_spk:
                ts   = ExportService._fmt_ts_simple(seg["start_time"])
                para = doc.add_paragraph()
                run  = para.add_run(f"SPEAKER {spk}  [{ts}]")
                run.bold = True
                run.font.color.rgb = color
                run.font.size = Pt(9)
                prev_spk = spk

            # Nội dung
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(16)

            if lang == "both":
                r_en = para.add_run(seg.get("text_en", "") + "\n")
                r_en.font.size = Pt(11)

                r_vi = para.add_run(seg.get("text_vi", ""))
                r_vi.font.size  = Pt(11)
                r_vi.font.color.rgb = RGBColor(0x6B, 0x77, 0x99)
            else:
                text = ExportService._get_text(seg, lang)
                run  = para.add_run(text)
                run.font.size = Pt(11)

        doc.save(path)
        print(f"[EXPORT] DOCX → {path} ({len(segments)} segments)")

    # ── Helpers ─────────────────────────────────────────────────────────────

    @staticmethod
    def _get_text(seg: dict, lang: str) -> str:
        if lang == "en":
            return seg.get("text_en", "")
        return seg.get("text_vi", "") or seg.get("text_en", "")

    @staticmethod
    def _fmt_srt_time(seconds: float) -> str:
        """00:00:01,234"""
        td  = timedelta(seconds=seconds)
        total_s = int(td.total_seconds())
        ms      = int((seconds - int(seconds)) * 1000)
        h, rem  = divmod(total_s, 3600)
        m, s    = divmod(rem, 60)
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    @staticmethod
    def _fmt_ts_simple(seconds: float) -> str:
        """0:01:23"""
        total = int(seconds)
        h, rem = divmod(total, 3600)
        m, s   = divmod(rem, 60)
        if h:
            return f"{h}:{m:02d}:{s:02d}"
        return f"{m}:{s:02d}"