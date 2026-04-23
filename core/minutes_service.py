import json
import os
import threading
from dataclasses import dataclass

import requests
from PyQt6.QtCore import QObject, pyqtSignal

# ── Groq config (giống summary_service.py) ───────────────────────────────────
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL   = "llama-3.3-70b-versatile"   # 70B tốt hơn 8B cho viết văn bản dài

# ── Prompt biên bản ───────────────────────────────────────────────────────────
def _load_prompt() -> str:
    path = os.path.join(
        os.path.dirname(__file__), "..", "dialogue_categories.json")
    try:
        with open(os.path.abspath(path), encoding="utf-8") as f:
            return json.load(f).get("Cuộc họp công việc", "")
    except Exception:
        return ""

MEETING_PROMPT = _load_prompt()

# Fallback nếu JSON không load được
_DEFAULT_PROMPT = """\
Từ nội dung hội thoại dưới đây, hãy tạo biên bản cuộc họp bằng tiếng Việt \
với định dạng Markdown sau. Chỉ trả về nội dung biên bản, không giải thích thêm.

# [Tên cuộc họp]

## Thông tin
- **Ngày họp:** ...
- **Người tham gia:** ...
- **Chủ đề:** ...

## Mục tiêu
...

## Tóm tắt
...

## Quyết định
- ...
- ...

## Bước tiếp theo
| Nội dung công việc | Thời hạn |
|---|---|
| ... | ... |
"""


# ── Data class ────────────────────────────────────────────────────────────────
@dataclass
class MinutesResult:
    text:  str = ""
    error: str = ""


# ── Service ───────────────────────────────────────────────────────────────────
class MinutesService:

    MAX_CHARS = 12_000

    def generate(self, segments: list[dict], title: str = "", custom_prompt: str = "", recorded_at: str = "") -> MinutesResult:
        """Tạo biên bản đồng bộ — gọi trong thread riêng, không block UI."""

        api_key = os.environ.get("GROQ_API_KEY", "").strip()
        if not api_key:
            return MinutesResult(error="no_api_key")

        transcript = self._build_transcript(segments)
        if not transcript.strip():
            return MinutesResult(error="empty_transcript")

        base = custom_prompt or MEETING_PROMPT or _DEFAULT_PROMPT
        meta = f"Tên cuộc họp: {title or 'Cuộc họp'}"
        if recorded_at:
            meta += f"\nNgày họp: {recorded_at}"
        user_msg = (
            f"{base}\n\n"
            f"{meta}\n\n"
            f"NỘI DUNG:\n{transcript}"
        )

        try:
            response = requests.post(
                GROQ_API_URL,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type":  "application/json",
                },
                json={
                    "model":       GROQ_MODEL,
                    "messages":    [{"role": "user", "content": user_msg}],
                    "max_tokens":  2000,
                    "temperature": 0.3,
                },
                timeout=30,
            )
            response.raise_for_status()
            text = response.json()["choices"][0]["message"]["content"]
            return MinutesResult(text=text.strip())

        except requests.Timeout:
            return MinutesResult(error="timeout")
        except requests.HTTPError as e:
            status = e.response.status_code if e.response else 0
            if status == 401:
                return MinutesResult(error="invalid_api_key")
            if status == 429:
                raise   # để worker retry
            return MinutesResult(error=f"http_{status}")
        except Exception as e:
            print(f"[MINUTES] Error: {e}")
            return MinutesResult(error=str(e))

    # ── Helper ────────────────────────────────────────────────────────────────
    def _build_transcript(self, segments: list[dict]) -> str:
        lines, last_spk = [], None
        for seg in segments:
            spk = seg.get("speaker_index", 1)
            txt = (seg.get("text_vi") or seg.get("text_en", "")).strip()
            if not txt:
                continue
            if spk != last_spk:
                lines.append(f"\nSpeaker {spk}:")
                last_spk = spk
            lines.append(txt)

        full = "\n".join(lines).strip()
        if len(full) > self.MAX_CHARS:
            full = full[:self.MAX_CHARS] + "\n[... transcript bị cắt bớt ...]"
        return full


# ── Qt Worker (non-blocking, có retry 429) ────────────────────────────────────
class MinutesWorker(QObject):
    """
    Chạy generate() trong thread riêng, emit signal khi xong.

    Usage:
        worker = MinutesWorker(segments, title="Sprint review")
        worker.done.connect(lambda r: show_minutes(r))
        worker.start()
    """
    done = pyqtSignal(object)   # MinutesResult

    def __init__(self, segments: list[dict], title: str = "",
                 custom_prompt: str = "", recorded_at: str = "", parent=None):
        super().__init__(parent)
        self._segments      = segments
        self._title         = title
        self._custom_prompt = custom_prompt
        self._recorded_at   = recorded_at
        self._svc           = MinutesService()

    def start(self):
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        import time
        for attempt in range(3):
            try:
                result = self._svc.generate(self._segments, self._title, self._custom_prompt, self._recorded_at)
                self.done.emit(result)
                return
            except Exception as e:
                msg = str(e)
                if "429" in msg and attempt < 2:
                    wait = (attempt + 1) * 5
                    print(f"[MINUTES] Rate limited, retry in {wait}s… (attempt {attempt+1}/3)")
                    time.sleep(wait)
                else:
                    self.done.emit(MinutesResult(error=msg))
                    return


# ── Export .docx ────────────────────────────────────────────────────────────────
def export_docx(minutes_text: str, title: str, save_path: str) -> str:
    """Export biên bản ra .docx. Trả về path đã lưu."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, hex_color: str):
        """Tô màu nền cell bảng."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _add_inline_runs(para, text: str, base_color: RGBColor | None = None):
        """Parse **bold** inline và thêm runs vào paragraph."""
        import re
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            else:
                run = para.add_run(part)
                if base_color:
                    run.font.color.rgb = base_color

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # Tiêu đề
    h = doc.add_heading(title or "Biên bản cuộc họp", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()

    # Thu thập các dòng bảng liên tiếp rồi xử lý cùng nhau
    lines = minutes_text.split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()

        # ── Bảng Markdown ──
        if s.startswith("|"):
            # Gom tất cả dòng bảng liên tiếp
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1

            # Tách header / separator / data rows
            rows = []
            for row in tbl_lines:
                cells = [c.strip() for c in row.strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    continue   # bỏ dòng separator
                rows.append(cells)

            if not rows:
                continue

            col_count = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=col_count)
            tbl.style = "Table Grid"

            for r_idx, row_cells in enumerate(rows):
                for c_idx, cell_text in enumerate(row_cells):
                    if c_idx >= col_count:
                        break
                    cell = tbl.cell(r_idx, c_idx)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _add_inline_runs(para, cell_text)
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                    if r_idx == 0:           # header row
                        _set_cell_bg(cell, "D9E1F2")
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            doc.add_paragraph()
            continue

        # ── Các loại dòng khác ──
        i += 1
        if not s:
            doc.add_paragraph()
        elif s.startswith("# "):
            ph = doc.add_heading(s[2:].strip(), level=1)
            if ph.runs:
                ph.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif s.startswith("## "):
            ph = doc.add_heading(s[3:].strip(), level=2)
            if ph.runs:
                ph.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif s.startswith("### "):
            ph = doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, s[2:].strip())
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
        elif len(s) > 2 and s[0].isdigit() and s[1] in ".)":
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, s[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, s)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)

    doc.save(save_path)
    return save_path